from pathlib import Path
import re
import ollama
import mmap
import numpy as np
import faiss

def _clean_text(s: str) -> str:
    # Collapse excessive whitespace, fix soft hyphenations, etc.
    s = s.replace('\u00ad', '')  # soft hyphen
    s = re.sub(r'[ \t]+', ' ', s)
    s = re.sub(r'\s+\n', '\n', s)
    return s.strip()

def load_pdf_text(path: str, password: str | None = None) -> list[str]:
    """Return a list of lines from a text-based PDF (no OCR)."""
    import fitz  # PyMuPDF
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(path)

    # If encrypted and password provided, PyMuPDF can authenticate directly 
    doc = fitz.open(p.as_posix())
    if doc.needs_pass:
        if not password:
            raise ValueError("PDF is encrypted and no password was provided.")
        if not doc.authenticate(password):
            raise ValueError("Invalid PDF password.")

    lines: list[str] = []
    for page in doc:
        text = page.get_text("text")  # “text” keeps natural reading order
        text = _clean_text(text)
        if text:
            # Keep page-break awareness if you like:
            lines.extend([ln for ln in text.splitlines() if ln.strip()])
    doc.close()
    return lines

def load_docx_text(path: str) -> list[str]:
    """Return a list of lines from a .docx file (no OCR)."""
    import docx  # python-docx
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(path)

    doc = docx.Document(p.as_posix())
    lines: list[str] = []

    # Paragraphs
    for para in doc.paragraphs:
        t = _clean_text(para.text)
        if t:
            lines.append(t)

    # Tables (cells often contain important text)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                t = _clean_text(cell.text)
                if t:
                    lines.append(t)

    return lines

def load_text_any(path: str, password: str | None = None) -> list[str]:
    """Dispatch loader by extension; returns list[str] lines ready for embedding."""
    ext = Path(path).suffix.lower()
    if ext == ".pdf":
        return load_pdf_text(path, password=password)
    elif ext == ".docx":
        return load_docx_text(path)
    elif ext in (".txt", ".md", ".csv"):
        return Path(path).read_text(encoding="utf-8", errors="replace").splitlines()
    else:
        raise ValueError(f"Unsupported file type: {ext} (supported: .pdf, .docx, .txt, .md, .csv)")


# --- Load dataset from a file path (PDF/Word/TXT) ---
SOURCE_PATH = r"assests/DINOv3.pdf"  
# dataset = load_text_any(SOURCE_PATH) 
def chunk_lines(lines, max_chars=1000):
    """Group lines into chunks of roughly `max_chars` characters."""
    chunks, current = [], ""
    for line in lines:
        if len(current) + len(line) < max_chars:
            current += line + " "
        else:
            chunks.append(current.strip())
            current = line + " "
    if current:
        chunks.append(current.strip())
    return chunks

raw_lines = load_text_any(SOURCE_PATH)
dataset = chunk_lines(raw_lines, max_chars=1000)
print(f"Created {len(dataset)} text chunks for embedding.")

print(f"Loaded {len(dataset)} entries from {SOURCE_PATH}")


# --- Define models ---
EMBEDDING_MODEL = 'hf.co/CompendiumLabs/bge-base-en-v1.5-gguf:F32'
LANGUAGE_MODEL  = 'hf.co/bartowski/Llama-3.2-1B-Instruct-GGUF:F16'

# --- Batch embed (single call instead of per line) ---
resp = ollama.embed(model=EMBEDDING_MODEL, input=dataset)
embeddings = resp['embeddings']  # list of embeddings, one per chunk
print(f'Generated {len(embeddings)} embeddings')

# --- Build NumPy arrays for fast retrieval ---
vecs  = np.array(embeddings, dtype=np.float32)
norms = np.linalg.norm(vecs, axis=1)


def build_faiss_index(vecs: np.ndarray):
    """Build and return a normalized FAISS index for cosine similarity."""
    vecs = vecs.astype('float32')
    # Normalize vectors (for cosine similarity)
    faiss.normalize_L2(vecs)
    
    d = vecs.shape[1]
    index = faiss.IndexFlatIP(d)  # IP = Inner Product (cosine similarity after normalization)
    index.add(vecs)
    return index

index = build_faiss_index(vecs)



# def retrieve(query, top_n=3):
#     q = np.array(ollama.embed(model=EMBEDDING_MODEL, input=query)['embeddings'][0], dtype=np.float32)
#     qn = np.linalg.norm(q)
#     sims = (vecs @ q) / (norms * (qn + 1e-12))
#     top_idx = np.argpartition(-sims, top_n)[:top_n]
#     top_idx = top_idx[np.argsort(-sims[top_idx])]
#     return [(dataset[i], float(sims[i])) for i in top_idx]

def retrieve(query: str, top_n: int = 3):
    """Retrieve top-N most similar entries using FAISS."""
    # Get query embedding
    q_emb = np.array(ollama.embed(model=EMBEDDING_MODEL, input=query)['embeddings'][0], dtype=np.float32)
    q_emb = q_emb.reshape(1, -1)
    faiss.normalize_L2(q_emb)

    # Search
    sims, idxs = index.search(q_emb, top_n)

    # Return [(text, score), ...]
    return [(dataset[i], float(sims[0][j])) for j, i in enumerate(idxs[0])]

# --- Chatbot loop ---
while True:
    print()
    input_query = input('Ask me a question: ')
    retrieved_knowledge = retrieve(input_query, top_n=4)

    print('Retrieved knowledge:')
    for chunk, similarity in retrieved_knowledge:
        print(f' - (similarity: {similarity:.2f}) {chunk}')

    instruction_prompt = (
        "You are a knowledgeable AI assistant.\n"
        "Use the context below to answer accurately and concisely.\n"
        "If the context doesn't contain the answer, say 'The provided documents do not contain enough information.'\n\n"
        "Context:\n" +
        "\n\n".join([f'### {i+1}. {chunk}' for i, (chunk, _) in enumerate(retrieved_knowledge)])
    )


    stream = ollama.chat(
        model=LANGUAGE_MODEL,
        messages=[
            {'role': 'system', 'content': instruction_prompt},
            {'role': 'user',   'content': input_query},
        ],
        stream=True,
    )

    print('Chatbot response:\n\n')
    for part in stream:
        print(part['message']['content'], end='', flush=True)
